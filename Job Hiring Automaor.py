from docx import Document
import re
import xlrd
import numpy
import sys
import os
import comtypes.client
import time

###INSTRUCTIONS###
'''
1 - FILL OUT THE FILE PATH TO THE APPLICATION FOLDER

2 - CREATE A FOLDER FOR RESUMES AND CONNECT THE FILE PATH TO THE FOLDER

3 - PUT UP TO 6 RESUMES IN THE FOLDER AND CONNECT THERE FILE PATHS

4 - CREATE A FOLDER FOR THE COVER LETTERS AND CONNECT THE FILE PATH

5 - PUT UP TO 5 COVER LETTERS IN THE FOLDER AND CONNECT THERE FILE PATHS

6 - CREATE A FOULDER FOR THE DATA AND CONNECT THE FILE PATH

7 - MAKE AN EXCEL WITH THE DATA
    - FIRST CELL IS COVER LETTER TYPE
    - SECOND CELL IS RESUME TYPE
    - THIRD CELL IS COMPANY NAME
    - FITH CELL IS COMPANY ADRESS
    - SIXTH CELL IS COMPANY AREA

8 - PUT THE EXCEL IN THE FOLDER AND CONNECT THE FILE PATH


'''

#Application package
app_pack = "C:/Users/Gianpaolo/Desktop/2A Job Applications/"

#Resume Types
resume_path = "03 - Resumes/"

resume_1 = "Automotive Resume (Style 3).docx"
resume_2 = "Manufacturing Resume (Style 3).docx"
resume_3 = "Product Development Resume (Style 3).docx"
resume_4 = "Robotics and Controls Resume (Style 3).docx"
resume_5 = "Software Resume (Style 3).docx"

#Cover Letter Types
template_path = "02 - Cover_Letters/"

template_name_1 = "Cover Letter Automotive.docx"
template_name_2 = "Cover Letter General.docx"
template_name_3 = "Cover Letter Manufacturing.docx"
template_name_4 = "Cover Letter Product Development.docx"
template_name_5 = "Cover Letter Robotics and Controls.docx"
template_name_6 = "Cover Letter Software.docx"

#Excel Path
source_path = "01 - Company_List/"

source_name = "company_list.xlsx"
sheet_name = 'Sheet1'

################################################################################

#PDF format type
wdFormatPDF = 17

#Company, Adress and Location
company_name = "x"
company_address = "y"
company_location = "z"

#Source data from excel into matrix
book = xlrd.open_workbook(app_pack+source_path+source_name)
sheet = book.sheet_by_name(sheet_name)
data = [[sheet.cell_value(r, c) for c in range(sheet.ncols)] for r in range(sheet.nrows)]
data_array = numpy.array(data)
y_data, x_data = data_array.shape

#Search algorithm
def docx_replace_regex(doc_obj, regex , replace):
    
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
 
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)
                

#PDF algorithm
def convert_to_pdf(x, y, wdFormatPDF):
    
    in_file = os.path.abspath(x)
    out_file = os.path.abspath(y)
    
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

def createFolder(app_path):
    try:
        if not os.path.exists(app_path):
            os.makedirs(app_path)
    except OSError:
        print ('Error: Creating directory: ' + directory)

for y in range(y_data):
    
    temp = int(float(data_array[y][0]))
    temp_2 = int(float(data_array[y][1]))
    
    #Resume Switch
    if (temp_2== 1):
        document_2 = resume_1
    elif (temp_2 == 2):
        document_2 = resume_2
    elif (temp_2 == 3):
        document_2 = resume_3
    elif (temp_2 == 4):
        document_2 = resume_4
    elif (temp_2 == 5):
        document_2 = resume_5
    else:
        print(data_array[y][2] + ": Resume Not Found")
    
    #Cover Letter Switch
    if (temp == 1):
        document = Document(app_pack+template_path+template_name_1)
    elif (temp == 2):
        document = Document(app_pack+template_path+template_name_2)
    elif (temp == 3):
        document = Document(app_pack+template_path+template_name_3)
    elif (temp == 4):
        document = Document(app_pack+template_path+template_name_4)
    elif (temp == 5):
        document = Document(app_pack+template_path+template_name_5)
    elif (temp == 6):
        document = Document(app_pack+template_path+template_name_6)
    else:
        print(data_array[y][2] + ": Template Not Found")
        continue
    
    #Save data from excel
    company_name = data_array[y][2]
    company_address = data_array[y][3]
    company_location = data_array[y][4]
    job_title = data_array[y][5]
    
    #Create application folder
    application_folder = app_pack + company_name + " Application"
    createFolder(application_folder)

    #Name file for word
    file_reference_word = application_folder + "/" + "Gianpaolo Pittis " + company_name + " Cover Letter.docx"
    
    #Name file for resume pdf
    file_reference_resume = app_pack + resume_path + document_2
   
    #Name file for PDF
    file_reference_pdf = application_folder + "/" + "Gianpaolo Pittis " + company_name + " Cover Letter.pdf" 
    
    #Name file for resume pdf
    file_reference_resume_pdf = application_folder + "/" + "Gianpaolo Pittis " + company_name + " Resume.pdf"
    
    #Replace job title
    regex4 = re.compile(r"jbxx")
    replace4 = job_title + "."
    
    #Replace company name 2
    regex5 = re.compile(r"xpx")
    
    #Replace company name
    regex1 = re.compile(r"xcom")
    replace1 = company_name
    
    #Replace address
    regex2 = re.compile(r"xdress")
    replace2 = company_address

    #Replace city/province
    regex3 = re.compile(r"xplace")
    replace3 = company_location
    
    #Replace function
    docx_replace_regex(document, regex1, replace1)
    docx_replace_regex(document, regex2, replace2)
    docx_replace_regex(document, regex3, replace3)
    docx_replace_regex(document, regex4, replace4)
    docx_replace_regex(document, regex5, replace1)
    
    #Save document as word
    document.save(file_reference_word)
    
    #Add Resume
    convert_to_pdf(file_reference_resume, file_reference_resume_pdf, wdFormatPDF)
 
    #Save document as PDF
    convert_to_pdf(file_reference_word, file_reference_pdf, wdFormatPDF)
    
    print("Done: " + company_name)

print("DONE")