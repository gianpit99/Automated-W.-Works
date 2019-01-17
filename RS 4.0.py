from docx import Document
import re
import xlrd
import numpy
import sys
import os
import comtypes.client
import time

###INSTRUCTIONS###
'''
1 - Create a folder on your desktop named "2A Job Applications"

2 - In that folder create 2 folders: 1. "02 - Cover_Letters" 2. " 03 - Resumes" 3. "01 - Cover_Letters"

3 - Place an excel document named "company_list"
        The first cell should be cover letter type
        The second cell should be resuem type
        The third cell shoulf be company name
        The fourth Cell should be company address
        The fith cell should be company region

4 - Connect all the cover letter template file paths to the code

5 - Connect all resume template file paths to the code


'''
#Application package
app_pack = "C:/Users/Gianpaolo\Desktop/2A Job Applications/"

#Resume Types
resume_1 = "Resume 6"
resume_2 = "Resume 6"
resume_3 = "Resume 6"
resume_4 = "Resume 6"
resume_5 = "Resume 6"

#Cover Letter Types
template_name_1 = "abc"
template_name_2 = "abc"
template_name_3 = "abc"
template_name_4 = "abc"
template_name_5 = "abc"

#PDF format type
wdFormatPDF = 17



#Name and destination of where you want the file saved
name_of_file = "Practice Path"
file_path = app_pack
file_reference = file_path+name_of_file+".docx"

#Template destinations
template_path = app_pack + "02 - Cover_letters/"

#Template Type 1

template_reference_1 = template_path+template_name_1+".docx"

#Template Type 2

template_reference_2 = template_path+template_name_2+".docx"

#Template Type 3

template_reference_3 = template_path+template_name_3+".docx"

#Template Type 4

template_reference_4 = template_path+template_name_4+".docx"

#Template Type 5

template_reference_5 = template_path+template_name_5+".docx"

#Name and destination of where the excel document is
source_name = "company_list"
source_path = app_pack + "01 - Company_list/"
source_reference = source_path+source_name+".xlsx"
sheet_name = 'Sheet1'

#Company, Adress and Location
company_name = "x"
company_address = "y"
company_location = "z"

#Source data from excel into matrix
book = xlrd.open_workbook(source_reference)
sheet = book.sheet_by_name('Sheet1')
data = [[sheet.cell_value(r, c) for c in range(sheet.ncols)] for r in range(sheet.nrows)]
data_array = numpy.array(data)
y_data, x_data = data_array.shape

#Search algorithm
def docx_replace_regex(doc_obj, regex , replace):
    
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
 
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)
                

#PDF algorithm
def convert_to_pdf(x, y, wdFormatPDF):
    
    in_file = os.path.abspath(x)
    out_file = os.path.abspath(y)
    
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

def createFolder(app_path):
    try:
        if not os.path.exists(app_path):
            os.makedirs(app_path)
    except OSError:
        print ('Error: Creating directory: ' + directory)

for y in range(y_data):
    
    temp = int(float(data_array[y][0]))
    
    #Cover Letter Switch
    if (temp == 1):
        document = Document(template_reference_1)
    elif (temp == 2):
        document = Document(template_reference_2)
    elif (temp == 3):
        document = Document(template_reference_3)
    elif (temp == 4):
        document = Document(template_reference_4)
    elif (temp == 5):
        document = Document(template_reference_5)
    else:
        print(data_array[y][2] + ": Template Not Found")
        continue
    
    #Save data from excel
    company_name = data_array[y][2]
    company_address = data_array[y][3]
    company_location = data_array[y][4]
    
    #Name file for word
    file_reference = file_path + company_name + " Cover Letter.docx"
   
    #Name file for PDF
    file_reference_pdf = file_path + company_name + "Cover Letter.pdf" 
    
    #Replace company name
    regex1 = re.compile(r"com")
    replace1 = company_name
    
    #Replace address
    regex2 = re.compile(r"dress")
    replace2 = company_address

    #Replace city/province
    regex3 = re.compile(r"place")
    replace3 = company_location
    
    #Replace function
    docx_replace_regex(document, regex1, replace1)
    docx_replace_regex(document, regex2, replace2)
    docx_replace_regex(document, regex3, replace3)
    
    #Create application folder
    application_folder = app_pack + company_name + " Application"
    print(application_folder)
    createFolder(application_folder)
    
    #Save document as word
    document.save(file_reference)
 
    time.sleep(2)
 
    #Save document as PDF
    convert_to_pdf(file_reference, file_reference_pdf, wdFormatPDF)